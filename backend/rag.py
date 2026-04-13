import os
import json
from dotenv import load_dotenv
load_dotenv()

from typing import TypedDict, Annotated, Sequence, Any, Optional, Literal
from pathlib import Path
import io
import re

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.retrievers import BM25Retriever
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage, ToolMessage
from langchain_core.tools import tool
from tools import get_search_knowledge_base_tool
from langchain_experimental.text_splitter import SemanticChunker
from pydantic import BaseModel, Field

from langgraph.prebuilt import create_react_agent
from langgraph.checkpoint.memory import MemorySaver

# ─── Firestore ────────────────────────────────────────────────────────────────

from firestore_service import (
    save_chunks_to_firestore,
    save_single_doc_to_firestore,
    load_all_chunks_from_firestore,
    delete_chunks_from_firestore,
    save_user_memory_firestore,
    load_user_memory_firestore,
)


# ─── User Memory Store ───────────────────────────────────────────────────────

def load_user_memory(user_id: str) -> list[str]:
    """Load stored facts for a given user from Firestore."""
    return load_user_memory_firestore(user_id)

def save_user_memory(user_id: str, facts: list[str]):
    """Overwrite stored facts for a given user in Firestore."""
    save_user_memory_firestore(user_id, facts)
    print(f"[Firestore] Saved {len(facts)} memory facts for user '{user_id}'")

# ─── System Prompt ────────────────────────────────────────────────────────────

def load_prompts(filepath: Path) -> dict[str, str]:
    if not filepath.exists():
        return {}
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    prompts = {}
    import re
    parts = re.split(r'^#\s+', content, flags=re.MULTILINE)
    for part in parts:
        if not part.strip(): continue
        lines = part.split('\n', 1)
        name = lines[0].strip()
        body = lines[1].strip() if len(lines) > 1 else ""
        prompts[name] = body
    return prompts

PROMPTS_FILE = Path(__file__).parent / "system_prompts.md"
PROMPTS = load_prompts(PROMPTS_FILE)
AGENT_SYSTEM_PROMPT = PROMPTS.get("AGENT_SYSTEM_PROMPT", "")

# ─── IT Support RAG Agent ────────────────────────────────────────────────────

class ITSupportRAG:
    def __init__(self):
        # LLM
        self.llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)

        # Retrieval mode — switchable per call for evaluation
        # "hybrid" = BM25 + Dense (RRF), "dense" = vector-only
        self._retrieval_mode = "hybrid"

        # Semantic chunker
        self.embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
        self.chunker = SemanticChunker(
            self.embeddings,
            breakpoint_threshold_type="percentile",
        )

        # Documents — load from Firestore
        self.documents = []
        firestore_chunks = load_all_chunks_from_firestore()
        for chunk_data in firestore_chunks:
            doc = Document(
                page_content=chunk_data["content"],
                metadata={
                    "source": chunk_data["source"],
                    "chunk": chunk_data.get("chunk_index", 0),
                    "total_chunks": chunk_data.get("total_chunks", 1),
                    "department": chunk_data.get("department", ""),
                },
            )
            self.documents.append(doc)
        print(f"[Firestore] Loaded {len(firestore_chunks)} chunks from Firestore")
        self._init_retrievers()

        # Memory extraction prompt
        memory_prompt_sys = PROMPTS.get("MEMORY_EXTRACT_PROMPT_SYSTEM", "")
        self.memory_extract_prompt = ChatPromptTemplate.from_messages([
            ("system", memory_prompt_sys),
            ("human", "User said: {user_message}\nAssistant replied: {assistant_reply}"),
        ])

        # Build the ReAct agent
        self.checkpointer = MemorySaver()
        self._build_agent()

    # ─── Retriever Setup ──────────────────────────────────────────────────

    def _init_retrievers(self):
        # ── BM25 (Sparse) ──────────────────────────────────────────────
        def bm25_preprocess(text: str) -> list[str]:
            text = text.lower()
            text = re.sub(r'[^\w\s]', ' ', text)  # strip punctuation
            return [t for t in text.split() if t]

        self.bm25_retriever = BM25Retriever.from_documents(
            self.documents,
            preprocess_func=bm25_preprocess,
        )
        self.bm25_retriever.k = 5

        # ── Chroma Vector Store (Dense) ────────────────────────────────
        # Clean up previous collection if it exists (for add/delete rebuilds)
        if hasattr(self, 'vectorstore'):
            try:
                self.vectorstore.delete_collection()
            except Exception:
                pass

        self.vectorstore = Chroma.from_documents(
            documents=self.documents,
            embedding=self.embeddings,
            collection_name="it_support",
        )
        self.dense_retriever = self.vectorstore.as_retriever(search_kwargs={"k": 5})

    # ─── Hybrid Retrieval (RRF) ──────────────────────────────────────────

    def _hybrid_retrieve(self, query: str, top_k: int = 3) -> list[Document]:
        """Combine BM25 + dense retrieval using Reciprocal Rank Fusion."""
        bm25_docs = self.bm25_retriever.invoke(query)
        dense_docs = self.dense_retriever.invoke(query)

        # Reciprocal Rank Fusion: score = sum(1 / (k + rank)) across retrievers
        RRF_K = 60
        doc_scores: dict[str, float] = {}
        doc_map: dict[str, Document] = {}

        for rank, doc in enumerate(bm25_docs):
            key = doc.page_content
            doc_scores[key] = doc_scores.get(key, 0) + 1.0 / (RRF_K + rank + 1)
            doc_map[key] = doc

        for rank, doc in enumerate(dense_docs):
            key = doc.page_content
            doc_scores[key] = doc_scores.get(key, 0) + 1.0 / (RRF_K + rank + 1)
            doc_map[key] = doc

        sorted_keys = sorted(doc_scores, key=doc_scores.get, reverse=True)
        return [doc_map[k] for k in sorted_keys[:top_k]]

    # ─── Build ReAct Agent ────────────────────────────────────────────────

    def _build_agent(self):
        """Build a ReAct agent with tools using LangGraph."""
        search_knowledge_base = get_search_knowledge_base_tool(self)

        self.tools = [search_knowledge_base]
        self._last_sources = []

        # Create the ReAct agent with LangGraph
        self.agent = create_react_agent(
            model=self.llm,
            tools=self.tools,
            checkpointer=self.checkpointer,
        )

    # ─── Source Filtering ─────────────────────────────────────────────────

    def _filter_cited_sources(self, answer: str) -> list[dict]:
        """Only return sources that the agent actually cited in its answer."""
        if not self._last_sources:
            return []

        cited = []
        for src in self._last_sources:
            source_id = src["source"]
            # Check if the source ID appears in the answer text
            if source_id in answer:
                cited.append(src)

        return cited

    # ─── Public API ───────────────────────────────────────────────────────

    def ask(self, question: str, user_id: str = "default", thread_id: str = "default") -> tuple[str, list, bool]:
        """Run the agent for a single question."""
        user_memory = load_user_memory(user_id)

        # Build system prompt with user memory
        if user_memory:
            memory_section = "Known facts about this user:\n" + "\n".join(f"- {f}" for f in user_memory)
        else:
            memory_section = ""

        system_prompt = AGENT_SYSTEM_PROMPT.format(user_memory_section=memory_section)

        config = {"configurable": {"thread_id": thread_id}}
        self._last_sources = []

        result = self.agent.invoke(
            {"messages": [SystemMessage(content=system_prompt), HumanMessage(content=question)]},
            config=config,
        )

        # Extract the final AI response
        final_message = result["messages"][-1]
        answer = final_message.content

        # Check if tools were used (= RAG was involved)
        is_from_rag = any(
            isinstance(m, ToolMessage) for m in result["messages"]
        )

        sources = self._filter_cited_sources(answer) if is_from_rag else []

        return answer, sources, is_from_rag

    async def ask_stream(self, question: str, user_id: str = "default", thread_id: str = "default"):
        """Async generator that yields SSE event dicts for streaming."""
        import asyncio

        user_memory = load_user_memory(user_id)

        # Build system prompt with user memory
        if user_memory:
            memory_section = "Known facts about this user:\n" + "\n".join(f"- {f}" for f in user_memory)
        else:
            memory_section = ""

        system_prompt = AGENT_SYSTEM_PROMPT.format(user_memory_section=memory_section)
        config = {"configurable": {"thread_id": thread_id}}
        self._last_sources = []

        input_messages = {
            "messages": [SystemMessage(content=system_prompt), HumanMessage(content=question)]
        }

        # Stream through the agent's ReAct loop
        full_answer = ""
        is_from_rag = False

        async for event in self.agent.astream_events(input_messages, config=config, version="v2"):
            kind = event.get("event", "")

            # Tool calls — agent decided to search
            if kind == "on_tool_start":
                tool_input = event.get("data", {}).get("input", {})
                query = tool_input.get("query", "") if isinstance(tool_input, dict) else str(tool_input)
                yield {"type": "status", "content": f"Searching knowledge base: \"{query[:50]}\"..."}
                is_from_rag = True

            # Tool finished
            elif kind == "on_tool_end":
                yield {"type": "status", "content": "Analyzing results..."}

            # LLM streaming tokens (final response)
            elif kind == "on_chat_model_stream":
                chunk = event.get("data", {}).get("chunk")
                if chunk and hasattr(chunk, "content") and chunk.content:
                    # Only stream tokens from the final response, not tool-calling reasoning
                    # Check if this is a tool call chunk
                    if hasattr(chunk, "tool_calls") and chunk.tool_calls:
                        continue
                    if hasattr(chunk, "tool_call_chunks") and chunk.tool_call_chunks:
                        continue
                    full_answer += chunk.content
                    yield {"type": "token", "content": chunk.content}

        # Step: Extract and save user memory
        if user_id and full_answer:
            try:
                existing = user_memory
                mem_response = self.llm.invoke(
                    self.memory_extract_prompt.format_messages(
                        existing_facts=json.dumps(existing) if existing else "None",
                        user_message=question,
                        assistant_reply=full_answer,
                    )
                )
                new_facts = json.loads(mem_response.content)
                if isinstance(new_facts, list) and new_facts:
                    combined = list(set(existing + new_facts))
                    save_user_memory(user_id, combined)
            except (json.JSONDecodeError, Exception):
                pass

        sources = self._filter_cited_sources(full_answer) if is_from_rag else []
        yield {"type": "done", "sources": sources, "is_from_rag": is_from_rag}

    def get_chat_history(self, thread_id: str) -> list[dict]:
        """Get chat history for a specific thread."""
        config = {"configurable": {"thread_id": thread_id}}
        try:
            checkpoint = self.checkpointer.get(config)
            if checkpoint and "channel_values" in checkpoint:
                messages = checkpoint["channel_values"].get("messages", [])
                history = []
                for m in messages:
                    if isinstance(m, HumanMessage):
                        history.append({"role": "user", "content": m.content})
                    elif isinstance(m, AIMessage) and m.content:
                        history.append({"role": "assistant", "content": m.content})
                return history
        except Exception:
            pass
        return []

    # ─── Document Management ──────────────────────────────────────────────

    def add_document(self, source: str, content: str, department: str = ""):
        # Use semantic chunking for longer documents
        MIN_CHUNK_LENGTH = 300
        OVERLAP_CHARS = 100

        if len(content) > MIN_CHUNK_LENGTH:
            chunks = self.chunker.split_text(content)

            # Add overlap: prepend tail of previous chunk to current chunk
            overlapped_chunks = []
            for i, chunk in enumerate(chunks):
                if i > 0 and len(chunks[i - 1]) > OVERLAP_CHARS:
                    overlap = chunks[i - 1][-OVERLAP_CHARS:]
                    chunk = overlap + " " + chunk
                overlapped_chunks.append(chunk)

            for i, chunk in enumerate(overlapped_chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={"source": source, "chunk": i, "total_chunks": len(overlapped_chunks), "department": department},
                )
                self.documents.append(doc)

            # Sync to Firestore
            firestore_chunks = [
                {"content": c, "source": source, "chunk_index": i, "total_chunks": len(overlapped_chunks), "department": department}
                for i, c in enumerate(overlapped_chunks)
            ]
            save_chunks_to_firestore(source, firestore_chunks)
            print(f"[Firestore] Saved {len(overlapped_chunks)} chunks for '{source}'")
        else:
            doc = Document(page_content=content, metadata={"source": source, "department": department})
            self.documents.append(doc)

            # Sync to Firestore
            save_single_doc_to_firestore(source, content)
            print(f"[Firestore] Saved single doc '{source}'")

        self._init_retrievers()
        # Rebuild agent with updated retrievers
        self._build_agent()

    def delete_document(self, source: str):
        self.documents = [doc for doc in self.documents if doc.metadata.get("source") != source]
        if not self.documents:
            self.documents = [Document(page_content="", metadata={"source": "empty"})]

        # Sync to Firestore
        delete_chunks_from_firestore(source)
        print(f"[Firestore] Deleted chunks for '{source}'")

        self._init_retrievers()
        # Rebuild agent with updated retrievers
        self._build_agent()

    def get_all_documents(self):
        return [
            {
                "source": doc.metadata.get("source", "Unknown"),
                "content": doc.page_content,
                "chunk_index": doc.metadata.get("chunk"),
                "total_chunks": doc.metadata.get("total_chunks"),
                "department": doc.metadata.get("department", ""),
            }
            for doc in self.documents
        ]

    async def parse_uploaded_file(self, file) -> tuple[str, str]:
        filename = file.filename
        file_ext = Path(filename).suffix.lower()
        content_bytes = await file.read()

        if file_ext == '.pdf':
            return self._parse_pdf(content_bytes, filename)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_docx(content_bytes, filename)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.gif']:
            return self._parse_image(content_bytes, filename)
        elif file_ext in ['.md', '.txt']:
            return self._parse_text(content_bytes, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    def _parse_pdf(self, content_bytes: bytes, filename: str) -> tuple[str, str]:
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            source = filename
            return text_content.strip(), source
        except ImportError:
            raise RuntimeError("PyPDF2 not installed. Install with: pip install PyPDF2")

    def _parse_docx(self, content_bytes: bytes, filename: str) -> tuple[str, str]:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content_bytes))
            text_content = ""
            for para in doc.paragraphs:
                text_content += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
            source = filename
            return text_content.strip(), source
        except ImportError:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")

    def _parse_image(self, content_bytes: bytes, filename: str) -> tuple[str, str]:
        import base64
        from langchain_openai import ChatOpenAI

        image_base64 = base64.b64encode(content_bytes).decode('utf-8')
        filename_lower = filename.lower()
        if filename_lower.endswith('.png'):
            media_type = 'image/png'
        elif filename_lower.endswith('.gif'):
            media_type = 'image/gif'
        elif filename_lower.endswith(('.jpg', '.jpeg')):
            media_type = 'image/jpeg'
        else:
            media_type = 'image/jpeg'

        source = filename
        try:
            vision_llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
            message = HumanMessage(
                content=[
                    {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{image_base64}"}},
                    {"type": "text", "text": """Extract all conversation history, text, or information from this image.
                    Format the extracted content as clean markdown with:
                    - Use headers (#, ##, ##) for structure
                    - Use bullet points for lists
                    - Use code blocks (```) for code or technical content
                    - Preserve the original meaning and structure as much as possible
                    - If it's a conversation, format as dialogue with clear speaker labels

                    Return ONLY the markdown content, no explanations."""},
                                    ],
                                )
            response = vision_llm.invoke([message])
            text_content = response.content.strip()
            if not text_content:
                text_content = f"[Image: {filename}] - No extractable content found"
            return text_content, source
        except Exception as e:
            raise RuntimeError(f"Failed to extract content from image: {str(e)}")

    def _parse_text(self, content_bytes: bytes, filename: str) -> tuple[str, str]:
        try:
            text_content = content_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text_content = content_bytes.decode('latin-1')
        source = filename
        return text_content.strip(), source


# ─── Singleton ────────────────────────────────────────────────────────────────

rag_system = None

def get_rag_system():
    global rag_system
    if rag_system is None:
        rag_system = ITSupportRAG()
    return rag_system

# ─── Evaluation Wrapper ───────────────────────────────────────────────────────

def rag_answer(
    query: str,
    retrieval_mode: str = "hybrid",
    top_k_search: int = 10,
    top_k_select: int = 3,
    use_rerank: bool = False,
    verbose: bool = False,
) -> dict:
    """
    Wrapper for eval.py to test the LangGraph RAG Agent.
    Switches between 'dense' and 'hybrid' retrieval per call.
    Returns the exact dictionary format expected by eval.py's scorecard.
    """
    rag = get_rag_system()

    # ── Switch retrieval mode before calling the agent ──
    rag._retrieval_mode = retrieval_mode
    if verbose:
        print(f"[rag_answer] mode={retrieval_mode}")

    # We pass the query directly to the ReAct agent
    answer, sources, is_from_rag = rag.ask(query)

    chunks_used = []
    if hasattr(rag, "_last_sources") and rag._last_sources:
        for src in rag._last_sources:
            chunks_used.append({
                "text": src.get("text", src.get("content", "")),
                "metadata": src.get("metadata", {"source": src.get("source", "unknown")}),
                "score": 1.0
            })

    return {
        "query": query,
        "answer": answer,
        "sources": [s["source"] for s in sources] if sources else [],
        "chunks_used": chunks_used,
        "config": {
            "retrieval_mode": retrieval_mode,
            "top_k_search": top_k_search,
            "top_k_select": top_k_select,
            "use_rerank": use_rerank,
            "label": f"langgraph_{retrieval_mode}"
        }
    }
